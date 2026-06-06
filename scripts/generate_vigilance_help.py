"""Generate the Operational Vigilance Report HELP guide in Word (.docx) and Excel (.xlsx).

Content is authored from the live module behaviour (backend/vigilance + OperationalVigilance.js).
Run: python3 /app/scripts/generate_vigilance_help.py
Output: /app/docs/Vigilance_Report_Help.docx and /app/docs/Vigilance_Report_Help.xlsx
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

OUT_DIR = "/app/docs"
os.makedirs(OUT_DIR, exist_ok=True)

NAVY = RGBColor(0x0B, 0x1F, 0x3B)
NAVY_HEX = "0B1F3B"
EMERALD_HEX = "10B981"
LIGHT_HEX = "EAF0F6"
AMBER_HEX = "FEF3C7"
WHITE_HEX = "FFFFFF"

# ============================================================================
# Shared content
# ============================================================================
ACCESS_ROWS = [
    ("HR Admin (role: hr)", "Full access", "Sees ALL employees and ALL uploaders merged into one row per employee/day. Can View / Edit / Delete any uploader's observation."),
    ("System Admin / Office Admin", "Full access", "Same as HR Admin (admin merged view)."),
    ("Employee with Designation = 'Vigilance'", "Operational access", "Sees the base grid for everyone but can Add / Edit / Delete only their OWN observations (strict isolation)."),
    ("Any other Employee", "No access", "Sees an 'Access Restricted' screen. The Vigilance Report menu item is hidden."),
]

COLUMN_REF = [
    ("Name", "System (read-only)", "Employee full name. Pulled live from the employee record.", "Text"),
    ("Date", "System (read-only)", "The working day for the row.", "DD-Mon-YYYY (e.g. 06-Jun-2026)"),
    ("Email-id", "System (read-only)", "Employee official email.", "Text"),
    ("Team", "System (read-only)", "Employee team.", "Text"),
    ("Punch-In", "System (read-only)", "First biometric/attendance punch-in.", "12-hour clock (e.g. 09:30 AM)"),
    ("Punch-Out", "System (read-only)", "Last attendance punch-out.", "12-hour clock (e.g. 06:15 PM)"),
    ("Total Hours", "System (read-only)", "Attendance total working hours.", "HH:MM"),
    ("System Login", "Vigilance-entered", "Time the employee logged into the work system.", "Enter 24h (e.g. 09:00); stored/shown 12h (09:00 AM)"),
    ("System Logout", "Vigilance-entered", "Time the employee logged out.", "Enter 24h (e.g. 18:30); stored/shown 12h (06:30 PM)"),
    ("Total Research Hours", "Vigilance-entered", "Total productive research duration for the day.", "HH:MM or HH:MM:SS (e.g. 08:00 or 08:15:30)"),
    ("Total Break Hours", "Vigilance-entered", "Total break duration for the day.", "HH:MM or HH:MM:SS (e.g. 01:00)"),
    ("Break blocks (From / To / Total)", "Vigilance-entered", "Morning / Lunch / Evening break and any Extra-Break(N). From/To are clock times, Total is a duration.", "From/To = clock; Total = HH:MM or HH:MM:SS"),
]

DUR_VALID = ["08:00", "08:15:30", "1:05", "00:45", "10:30:45"]
DUR_INVALID = ["25:99", "12:70", "10::30", "10-30", "AA:BB"]

HR_STEPS = [
    ("Open the module", "Click 'Vigilance Report' in the admin sidebar. The page opens at /vigilance and ALWAYS shows the base grid (every active employee x each day in the selected range), even before any sheet is uploaded."),
    ("Set filters", "Use Employee Name, From Date, To Date, Department, Designation and Team. Click 'Filter' to apply. Default range is Today -> Today."),
    ("Download Sample Template", "Click 'Download Sample Template'. The .xlsx is pre-filled with the active employees x days that match your CURRENT filters, with attendance columns already populated. Both From Date and To Date are required."),
    ("Distribute / collect", "Send the template to the vigilance team, or have them use the in-app upload. Each uploader fills only their observations."),
    ("Upload a filled sheet", "Click 'Upload Filled Sheet' and select the .xlsx. The system validates every row all-or-nothing; on any error nothing is saved and row-level messages are shown."),
    ("Review the merged grid", "Admin view shows ONE row per employee/day with a column group per uploader. Use the sticky right 'Actions' column to View / Edit / Delete a specific uploader's entry."),
    ("Sort, paginate, export", "Click any column header to sort (asc -> desc -> reset). Choose rows-per-page (10/25/50/100). Click 'Export' to download the current filtered dataset as .xlsx."),
]

USER_STEPS = [
    ("Open the module", "Vigilance-designation employees see 'Vigilance Report' in the employee sidebar (/employee/vigilance). You can open it even while onboarding is pending."),
    ("See the base grid", "You see every active employee x day for the selected range (system columns are read-only). Your own observations merge on top by employee + date."),
    ("Add an observation", "On a row, click the pencil ('Add observation') to open the dialog. Fill System Login/Logout, Total Research Hours, Total Break Hours and add break blocks, then click 'Save Entry'."),
    ("Use Download Template / Upload", "Optionally download the filter-aware template, fill it offline, and upload via 'Upload Filled Sheet'. Your upload only affects YOUR own observations."),
    ("Edit / Delete your entry", "After saving, the pencil edits your entry and the red trash deletes it. The delete button appears ONLY after an entry exists. You can View (read-only) any time."),
    ("Isolation", "You can see and change ONLY your own entries. You can never see or modify another vigilance user's observations."),
]

FAQ = [
    ("Why do I see 'Access Restricted'?", "Your account is not an admin and your designation is not 'Vigilance'. Ask HR to set your designation, or use an authorised account."),
    ("The table is empty / shows only dashes.", "No vigilance data has been entered for the selected range yet. The base grid still lists employees; '-' means no observation/attendance value for that cell."),
    ("My upload was rejected.", "Validation is all-or-nothing. Fix every row-level error shown (invalid time/duration, etc.) and re-upload. Nothing is saved until the whole sheet is valid."),
    ("What time format do I enter?", "Clock fields (System Login/Logout, break From/To) accept 24-hour (e.g. 13:45) and are stored/shown as 12-hour (01:45 PM). Duration fields accept HH:MM or HH:MM:SS."),
    ("Can two people fill the same employee/day?", "Yes. Each uploader has their own observation. Admin sees them merged side-by-side per uploader; each user only sees their own."),
    ("How are Extra breaks handled?", "Add as many break blocks as needed (Morning/Lunch/Evening + Extra-Break1..N). Columns are created dynamically and export/import them automatically."),
    ("Why is the Date shown as 06-Jun-2026?", "The whole HRMS uses the DD-Mon-YYYY format for consistency. Times of day keep AM/PM."),
    ("Download Template button is disabled.", "Select BOTH a From Date and a To Date. The validation message is 'Please select both From Date and To Date'."),
]

# ============================================================================
# WORD
# ============================================================================
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=9.5)
        shade_cell(hdr[i], NAVY_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9.5)
            if r_idx % 2 == 1:
                shade_cell(cells[i], LIGHT_HEX)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    p.space_after = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    return p


def body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def build_word():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BluBridge HRMS")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Operational Vigilance Report\nUser Help Guide")
    r2.bold = True
    r2.font.size = Pt(26)
    r2.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("For HR / Admin and Vigilance Users")
    rs.font.size = Pt(12)
    rs.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version 1.0  •  Updated 06-Jun-2026").font.size = Pt(9)
    doc.add_paragraph()

    # 1. Overview
    h1(doc, "1. Overview")
    body(doc, "The Operational Vigilance Report records daily productivity observations (system login/logout, "
              "research hours, break hours and individual break blocks) on top of each employee's attendance. "
              "The table ALWAYS shows a base grid of active employees for the selected date range; vigilance "
              "observations are merged on top by employee and date. System columns come live from Attendance and "
              "are never edited here.")

    # 2. Who can access
    h1(doc, "2. Who Can Access What")
    add_table(doc, ["Role", "Access", "What they can do"], ACCESS_ROWS,
              col_widths=[2.0, 1.2, 3.3])

    # 3. Screen layout
    doc.add_page_break()
    h1(doc, "3. Screen Layout & Controls")
    body(doc, "Filters row: Employee Name, From Date, To Date, Department, Designation, Team. "
              "Action buttons: Filter, Download Sample Template, Upload Filled Sheet, Add Entry, Export.")
    body(doc, "Table: sticky Name and Date columns on the left, a sticky Actions column on the right, a 2-row "
              "grouped header that stays pinned while scrolling, a synchronised top + bottom horizontal scrollbar, "
              "column sorting on every header, and pagination (10/25/50/100 rows per page).")

    # 4. Column reference
    h1(doc, "4. Column Reference")
    add_table(doc, ["Column", "Source", "Meaning", "Format"], COLUMN_REF,
              col_widths=[1.6, 1.3, 2.4, 1.9])

    # 5. Time & duration rules
    doc.add_page_break()
    h1(doc, "5. Time & Duration Rules")
    h2(doc, "Clock fields (System Login/Logout, break From/To)")
    body(doc, "Enter in 24-hour format (e.g. 13:45, 18:00, 00:15). Stored and displayed as 12-hour AM/PM "
              "(13:45 -> 01:45 PM, 00:15 -> 12:15 AM).")
    h2(doc, "Duration fields (Total Research Hours, Total Break Hours, each break Total)")
    body(doc, "Accept BOTH HH:MM and HH:MM:SS. Stored canonically and shown clean: HH:MM when seconds are :00, "
              "HH:MM:SS when seconds are present.")
    add_table(doc, ["Accepted (valid)", "Rejected (invalid)"],
              list(zip(DUR_VALID, DUR_INVALID + [""] * (len(DUR_VALID) - len(DUR_INVALID)))),
              col_widths=[2.6, 2.6])

    # 6. HR workflow
    h1(doc, "6. HR / Admin — Step by Step")
    add_table(doc, ["Step", "Action"], [(s[0], s[1]) for s in HR_STEPS], col_widths=[1.8, 4.7])

    # 7. User workflow
    doc.add_page_break()
    h1(doc, "7. Vigilance User — Step by Step")
    add_table(doc, ["Step", "Action"], [(s[0], s[1]) for s in USER_STEPS], col_widths=[1.8, 4.7])

    # 8. Data visibility & isolation
    h1(doc, "8. Data Visibility & Isolation Rules")
    for line in [
        "Admin sees every employee and every uploader, merged into one row per employee/day.",
        "A vigilance user sees the base grid for everyone but can add/edit/delete ONLY their own observations.",
        "A vigilance user can never see or modify another user's observations (enforced by the server — cross-user edit/delete returns a permission error).",
        "System columns (Name, Email, Team, Punch-In/Out, Total Hours) are read-only and come from Attendance.",
        "Deleting a vigilance row removes only that observation — employee, attendance and other HRMS data are untouched.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line).font.size = Pt(10.5)

    # 9. Upload validation
    h1(doc, "9. Upload Validation")
    for line in [
        "Validation is ALL-OR-NOTHING: if any row is invalid, nothing is saved and you get row-level error messages.",
        "Keep the header rows intact. Break groups (Morning/Lunch/Evening and Extra-Break(N)) are detected automatically.",
        "Unlimited Extra-Break columns are supported — add Extra-Break3, Extra-Break4, ... as needed.",
        "Re-uploading for the same employee/day updates your existing observation (upsert).",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line).font.size = Pt(10.5)

    # 10. FAQ
    doc.add_page_break()
    h1(doc, "10. Frequently Asked Questions")
    add_table(doc, ["Question", "Answer"], FAQ, col_widths=[2.3, 4.2])

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run("BluBridge HRMS — Operational Vigilance Report Help Guide • Confidential").font.size = Pt(8)

    path = os.path.join(OUT_DIR, "Vigilance_Report_Help.docx")
    doc.save(path)
    return path


# ============================================================================
# EXCEL
# ============================================================================
HDR_FILL = PatternFill("solid", fgColor=NAVY_HEX)
HDR_FONT = Font(bold=True, color=WHITE_HEX, size=11)
TITLE_FONT = Font(bold=True, color=NAVY_HEX, size=16)
SUB_FONT = Font(italic=True, color="444444", size=10)
ALT_FILL = PatternFill("solid", fgColor=LIGHT_HEX)
WRAP = Alignment(wrap_text=True, vertical="top")
WRAP_CTR = Alignment(wrap_text=True, vertical="center", horizontal="center")
THIN = Side(style="thin", color="CCCCCC")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_sheet(ws, title, subtitle, headers, rows, widths):
    ws.sheet_view.showGridLines = False
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    c = ws.cell(row=1, column=1, value=title)
    c.font = TITLE_FONT
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    s = ws.cell(row=2, column=1, value=subtitle)
    s.font = SUB_FONT
    ws.row_dimensions[1].height = 24

    hr = 4
    for j, h in enumerate(headers, start=1):
        cell = ws.cell(row=hr, column=j, value=h)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = WRAP_CTR
        cell.border = BORDER
    ws.row_dimensions[hr].height = 22

    for i, row in enumerate(rows):
        r = hr + 1 + i
        for j, val in enumerate(row, start=1):
            cell = ws.cell(row=r, column=j, value=val)
            cell.alignment = WRAP
            cell.border = BORDER
            cell.font = Font(size=10)
            if i % 2 == 1:
                cell.fill = ALT_FILL
    for j, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(j)].width = w
    ws.freeze_panes = ws.cell(row=hr + 1, column=1)


def build_excel():
    wb = Workbook()

    # Sheet: Overview / Access
    ws = wb.active
    ws.title = "Access & Roles"
    style_sheet(
        ws, "Operational Vigilance Report — Help",
        "Access matrix • Version 1.0 • Updated 06-Jun-2026",
        ["Role", "Access", "What they can do"],
        ACCESS_ROWS, [34, 16, 70],
    )

    # Sheet: Column reference
    ws2 = wb.create_sheet("Column Reference")
    style_sheet(
        ws2, "Column Reference",
        "System columns are read-only (from Attendance). Vigilance columns are entered/editable.",
        ["Column", "Source", "Meaning", "Format"],
        COLUMN_REF, [26, 20, 50, 40],
    )

    # Sheet: Time & duration rules
    ws3 = wb.create_sheet("Time & Duration Rules")
    dur_rows = []
    for i in range(max(len(DUR_VALID), len(DUR_INVALID))):
        v = DUR_VALID[i] if i < len(DUR_VALID) else ""
        x = DUR_INVALID[i] if i < len(DUR_INVALID) else ""
        dur_rows.append((v, x))
    rules = [
        ("Clock fields", "System Login/Logout, break From/To", "Enter 24-hour (13:45)", "Shown 12-hour (01:45 PM)"),
        ("Duration fields", "Total Research/Break Hours, break Total", "Enter HH:MM or HH:MM:SS", "Shown clean (HH:MM, or HH:MM:SS if seconds)"),
        ("Date", "Date column / pickers", "Pick from calendar", "Shown DD-Mon-YYYY (06-Jun-2026)"),
    ]
    style_sheet(
        ws3, "Time & Duration Rules",
        "How to enter clock times, durations and dates.",
        ["Field type", "Applies to", "Enter as", "Stored / shown as"],
        rules, [16, 38, 26, 40],
    )
    # append valid/invalid table below
    start = 4 + 1 + len(rules) + 2
    ws3.cell(row=start, column=1, value="Duration examples").font = Font(bold=True, color=NAVY_HEX, size=12)
    hr = start + 1
    for j, h in enumerate(["Accepted (valid)", "Rejected (invalid)"], start=1):
        cell = ws3.cell(row=hr, column=j, value=h)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = WRAP_CTR
        cell.border = BORDER
    for i, (v, x) in enumerate(dur_rows):
        r = hr + 1 + i
        cv = ws3.cell(row=r, column=1, value=v)
        cx = ws3.cell(row=r, column=2, value=x)
        for cell in (cv, cx):
            cell.border = BORDER
            cell.font = Font(size=10)
            cell.alignment = WRAP
        cv.fill = PatternFill("solid", fgColor="E7F6EC")
        if x:
            cx.fill = PatternFill("solid", fgColor="FDECEC")

    # Sheet: HR steps
    ws4 = wb.create_sheet("HR Admin Steps")
    style_sheet(
        ws4, "HR / Admin — Step by Step",
        "End-to-end workflow for HR / System Admin / Office Admin.",
        ["#", "Step", "Action"],
        [(i + 1, s[0], s[1]) for i, s in enumerate(HR_STEPS)], [5, 26, 80],
    )

    # Sheet: User steps
    ws5 = wb.create_sheet("Vigilance User Steps")
    style_sheet(
        ws5, "Vigilance User — Step by Step",
        "Workflow for employees with Designation = 'Vigilance'.",
        ["#", "Step", "Action"],
        [(i + 1, s[0], s[1]) for i, s in enumerate(USER_STEPS)], [5, 26, 80],
    )

    # Sheet: FAQ
    ws6 = wb.create_sheet("FAQ")
    style_sheet(
        ws6, "Frequently Asked Questions",
        "Common questions & answers.",
        ["Question", "Answer"],
        FAQ, [42, 78],
    )

    path = os.path.join(OUT_DIR, "Vigilance_Report_Help.xlsx")
    wb.save(path)
    return path


if __name__ == "__main__":
    w = build_word()
    x = build_excel()
    print("WORD :", w, os.path.getsize(w), "bytes")
    print("EXCEL:", x, os.path.getsize(x), "bytes")
