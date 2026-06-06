"""Generate the Vigilance Module USER GUIDE (for Designation = 'Vigilance' employees).

Audience: a vigilance employee with ZERO prior knowledge of the module.
Scope:    ONLY the vigilance-employee workflow — NO admin workflows/views/permissions.

Outputs:
  /app/docs/Vigilance_Module_User_Guide.pdf   (primary)
  /app/docs/Vigilance_Module_User_Guide.docx  (fallback)

Run: python3 /app/backend/scripts/generate_vigilance_user_guide.py
"""
import os

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle,
    PageBreak, ListFlowable, ListItem, NextPageTemplate,
)
from reportlab.platypus.tableofcontents import TableOfContents

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "/app/docs"
os.makedirs(OUT_DIR, exist_ok=True)

NAVY = colors.HexColor("#0B1F3B")
ACCENT = colors.HexColor("#1D4ED8")
EMERALD = colors.HexColor("#10B981")
RED = colors.HexColor("#DC2626")
LIGHT = colors.HexColor("#EAF0F6")
GREY = colors.HexColor("#64748B")
ROW_ALT = colors.HexColor("#F4F7FB")

NAVY_RGB = RGBColor(0x0B, 0x1F, 0x3B)
NAVY_HEX = "0B1F3B"
LIGHT_HEX = "EAF0F6"

TITLE = "Vigilance Module"
SUBTITLE = "User Guide for Vigilance Team Members"
VERSION = "Version 1.0  •  Updated 06-Jun-2026"
BRAND = "BluBridge HRMS"

# ===========================================================================
# CONTENT (shared by PDF + DOCX)
# Each section: ("Section Title", [blocks])
# block kinds:
#   ("p", text)                         -> paragraph
#   ("h", text)                         -> sub-heading
#   ("steps", [str, ...])               -> numbered list
#   ("bullets", [str, ...])             -> bullet list
#   ("note", text)                      -> highlighted note box
#   ("table", [headers], [rows])        -> table
# ===========================================================================
COLUMN_TABLE_HEAD = ["Column", "Type", "Who fills it", "What to enter"]
COLUMN_TABLE_ROWS = [
    ["Name", "System (locked)", "Auto", "Do not change — employee full name."],
    ["Email-id", "System (locked)", "Auto", "Do not change — official email."],
    ["Team", "System (locked)", "Auto", "Do not change — employee team."],
    ["Date", "System (locked)", "Auto", "Do not change — the working day (DD-Mon-YYYY)."],
    ["Punch-In", "System (locked)", "Auto", "Do not change — first attendance punch (12-hour)."],
    ["Punch-Out", "System (locked)", "Auto", "Do not change — last attendance punch (12-hour)."],
    ["Total Hours", "System (locked)", "Auto", "Do not change — attendance total (HH:MM)."],
    ["System Login", "Editable", "You", "Time the employee logged into the work system (24-hour, e.g. 09:30)."],
    ["System Logout", "Editable", "You", "Time the employee logged out (24-hour, e.g. 18:00)."],
    ["Total Research Hours", "Editable", "You", "Total productive research time (HH:MM or HH:MM:SS)."],
    ["Total Break Hours", "Editable", "You", "Total break time for the day (HH:MM or HH:MM:SS)."],
    ["Morning Break (From/To/Total)", "Editable", "You", "From/To are clock times; Total is a duration."],
    ["Lunch Break (From/To/Total)", "Editable", "You", "From/To are clock times; Total is a duration."],
    ["Evening Break (From/To/Total)", "Editable", "You", "From/To are clock times; Total is a duration."],
    ["Extra-Break1, 2, 3 ... N", "Editable", "You", "Add as many extra breaks as needed — columns are created automatically."],
]

TIME_TABLE_HEAD = ["Field", "Enter as", "Example", "Shown in HRMS as"]
TIME_TABLE_ROWS = [
    ["System Login / Logout", "24-hour", "18:00", "06:00 PM"],
    ["Break From / To times", "24-hour", "13:45", "01:45 PM"],
    ["Punch-In / Punch-Out (template)", "12-hour (already filled)", "09:30 AM", "09:30 AM"],
    ["Total Hours (template)", "HH:MM (already filled)", "08:30", "08:30"],
    ["Total Research / Break Hours", "HH:MM or HH:MM:SS", "08:15 or 08:15:30", "08:15 / 08:15:30"],
    ["Each Break Total", "HH:MM or HH:MM:SS", "00:45 or 00:45:20", "00:45 / 00:45:20"],
]

VALID_TABLE_HEAD = ["Valid duration", "Valid time (24h)", "Invalid (will be rejected)"]
VALID_TABLE_ROWS = [
    ["08:00", "09:30", "25:99  (hour/min out of range)"],
    ["08:15:30", "18:00", "12:70  (minutes > 59)"],
    ["00:45", "13:45", "6:00 PM  (don't type AM/PM)"],
    ["10:30:45", "00:15", "10-30  (use a colon, not a dash)"],
    ["1:05", "23:59", "blank in a required field"],
]

FAQ = [
    ["My upload failed / was rejected.",
     "Validation is all-or-nothing: if even one row has an error, nothing is saved. Read the row-level error messages, fix every listed row (usually a wrong time or duration format), then upload again."],
    ["What time format do I use?",
     "Type System Login/Logout and break From/To in 24-hour format (e.g. 18:00). The HRMS shows them back as 12-hour (06:00 PM). Durations accept HH:MM or HH:MM:SS."],
    ["A duration like '12:70' is refused.",
     "Minutes and seconds must be 00–59 and hours valid. Use a colon (:) as the separator and never add AM/PM to duration fields."],
    ["An employee is missing from my template.",
     "The template only includes ACTIVE employees who fall inside the From–To date range and match your selected filters. Widen the date range or clear the Department/Team/Designation filters and download again."],
    ["I can't see data I uploaded.",
     "Check that the From/To dates in the filters cover the dates you uploaded, then click Filter. Also remember you only ever see YOUR OWN observations — never another vigilance member's."],
    ["Download Sample Template button seems disabled.",
     "You must select BOTH a From Date and a To Date. The message 'Please select both From Date and To Date' will appear until you do."],
    ["I changed a system column (Name/Email/Date) by mistake.",
     "Re-download a fresh template and copy your entered values into it. Never edit the locked system columns — the upload matches rows by employee + date."],
    ["Why is the date shown like 06-Jun-2026?",
     "The whole HRMS uses the DD-Mon-YYYY date format for clarity. Times of day still use AM/PM."],
]

SECTIONS = [
    ("1. Introduction", [
        ("p", "Welcome! This guide is written for <b>Vigilance Team members</b>. It assumes you have never opened the Vigilance Module before and walks you through everything, step by step, so you can work independently without any training."),
        ("h", "What is the Vigilance Module?"),
        ("p", "The Vigilance Module is the part of the HRMS where you record daily <b>productivity observations</b> for employees — their system login/logout times, total research hours, total break time, and each individual break (morning, lunch, evening and any extra breaks)."),
        ("h", "Why does it exist?"),
        ("p", "It sits on top of the attendance system and gives the organisation a clear, day-by-day picture of how working time is actually spent. This helps ensure fairness, transparency and healthy work patterns."),
        ("h", "What are you expected to do?"),
        ("bullets", [
            "Download a ready-made template for the employees and dates you are responsible for.",
            "Fill in the productivity/observation values (the system fills the attendance values for you).",
            "Upload the completed sheet back into the module — or type entries directly in the app.",
            "Review, edit and export your own records whenever needed.",
        ]),
        ("note", "Important: You only ever see and manage YOUR OWN observations. Other vigilance members cannot see your data, and you cannot see theirs."),
    ]),
    ("2. Accessing the Module", [
        ("h", "Step-by-step"),
        ("steps", [
            "Open the HRMS in your browser and log in with your username and password.",
            "On the left sidebar, click <b>Vigilance Report</b>. The module opens.",
            "Take a moment to understand the screen layout (described below).",
        ]),
        ("h", "What you see on the screen"),
        ("bullets", [
            "<b>Filters row</b> — Employee Name, From Date, To Date, Department, Designation, Team, and a Filter button.",
            "<b>Action buttons</b> — Download Sample Template, Upload Filled Sheet, Add Entry, Export, and Download User Guide.",
            "<b>Data table</b> — a base grid that always lists employees for the selected dates. Your observations appear merged on top, by employee and date.",
        ]),
        ("note", "The table is always visible — even before you upload anything. Empty cells simply show a dash (-)."),
    ]),
    ("3. Understanding Filters", [
        ("p", "Filters let you narrow the table and the downloaded template to exactly the employees and dates you need. <b>Filters apply ONLY after you click the Filter button.</b>"),
        ("table", ["Filter", "What it does / when to use it"], [
            ["Employee Name", "Type to search for one specific person. Use when you only need a single employee."],
            ["From Date", "The first day of the range you want. Required for downloading a template."],
            ["To Date", "The last day of the range. Required for downloading a template."],
            ["Department", "Limit to one department. Use to focus on a group of teams."],
            ["Designation", "Limit to one job role across the company."],
            ["Team", "Limit to a single team. The most common way to focus your work."],
            ["Filter button", "Applies all the choices above. Nothing changes until you click it."],
        ]),
        ("h", "Best practices"),
        ("bullets", [
            "Always set the From and To dates first.",
            "Add a Team or Department filter to keep the list short and manageable.",
            "Click Filter after every change — otherwise the old view stays.",
        ]),
        ("h", "Example"),
        ("p", "To work on the 'Vigilance' team for the first week of June: set From = 01-Jun-2026, To = 07-Jun-2026, choose Team = Vigilance, then click Filter."),
    ]),
    ("4. Downloading the Sample Template", [
        ("h", "Complete workflow"),
        ("steps", [
            "Select the <b>From Date</b> and <b>To Date</b> (both are required).",
            "Optionally narrow the list with Department, Team, Employee or Designation filters and click Filter.",
            "Click <b>Download Sample Template</b>. An Excel (.xlsx) file downloads to your computer.",
        ]),
        ("h", "What the template already contains"),
        ("p", "The system automatically pre-fills the attendance/identity columns so you never have to type them:"),
        ("bullets", [
            "Employee Name", "Email ID", "Team", "Date",
            "Punch-In", "Punch-Out", "Total Hours",
        ]),
        ("h", "How rows are generated"),
        ("bullets", [
            "Only <b>active</b> employees inside the selected date range and matching filters are included.",
            "If From Date and To Date are the <b>same day</b>, each employee gets exactly one row.",
            "If you choose a <b>range of days</b>, each employee gets one row per applicable day.",
        ]),
        ("note", "Example: 3 employees × 5 days = 15 rows, each pre-filled with that day's attendance."),
    ]),
    ("5. Understanding Template Columns", [
        ("p", "The template has two kinds of columns. Learn the difference before you start typing."),
        ("h", "System columns — locked, do NOT change"),
        ("p", "These are generated for you from attendance and employee records. Editing them can break the upload, which matches rows by employee + date."),
        ("h", "Editable columns — your observations"),
        ("p", "These are the values you fill in. The full list is below."),
        ("table", COLUMN_TABLE_HEAD, COLUMN_TABLE_ROWS),
        ("note", "Extra-Break3, Extra-Break4, Extra-Break5 and beyond are fully supported — just add the columns/values and they are detected automatically."),
    ]),
    ("6. Time Format Rules", [
        ("p", "Getting the format right is the single most important thing for a successful upload."),
        ("h", "Clock times"),
        ("bullets", [
            "Type System Login, System Logout and break From/To in <b>24-hour</b> format (e.g. 09:30, 13:45, 18:00).",
            "After upload, the HRMS shows them in <b>12-hour</b> format automatically (18:00 becomes 06:00 PM).",
            "In the downloaded template, <b>Punch-In, Punch-Out and Total Hours stay in 12-hour / HH:MM</b> — they are already filled, leave them as-is.",
        ]),
        ("h", "Durations"),
        ("p", "Duration fields (Total Research Hours, Total Break Hours, each Break Total) accept BOTH <b>HH:MM</b> and <b>HH:MM:SS</b>."),
        ("table", TIME_TABLE_HEAD, TIME_TABLE_ROWS),
        ("h", "Accepted vs rejected examples"),
        ("table", VALID_TABLE_HEAD, VALID_TABLE_ROWS),
    ]),
    ("7. Filling the Template Correctly", [
        ("h", "Do"),
        ("bullets", [
            "Use 24-hour time for login/logout and break From/To.",
            "Use a colon (:) as the separator everywhere (09:30, 00:45:20).",
            "Leave the locked system columns exactly as downloaded.",
            "Double-check minutes/seconds are between 00 and 59.",
        ]),
        ("h", "Don't"),
        ("bullets", [
            "Don't add AM/PM to the fields you type.",
            "Don't use dashes or dots as separators (10-30, 10.30 are invalid).",
            "Don't rename, reorder or delete columns.",
            "Don't edit Name, Email, Team, Date, Punch-In, Punch-Out or Total Hours.",
        ]),
        ("h", "Common mistakes"),
        ("bullets", [
            "Wrong time format (e.g. '6 PM' instead of 18:00).",
            "Out-of-range values (e.g. 12:70).",
            "Uploading a non-Excel file or a renamed template.",
            "Editing a system-generated column.",
        ]),
    ]),
    ("8. Uploading the Filled Sheet", [
        ("h", "Step-by-step"),
        ("steps", [
            "Open the Vigilance Module.",
            "Click <b>Upload Filled Sheet</b>.",
            "Select your completed Excel file.",
            "Confirm the upload and wait for the result message.",
        ]),
        ("h", "What happens"),
        ("bullets", [
            "Every row is validated <b>all-or-nothing</b>. If any row is invalid, nothing is saved and you get clear, row-by-row error messages.",
            "When all rows are valid, your data is saved securely and appears in the table immediately.",
            "Uploading again for the same employee + date <b>updates</b> your existing entry (it does not create duplicates).",
        ]),
        ("note", "Privacy: your uploaded data is private to you. Other vigilance members cannot see or change it."),
    ]),
    ("9. Viewing Uploaded Data", [
        ("p", "Your records appear in the table, merged onto the base grid by employee and date."),
        ("bullets", [
            "<b>Sorting</b> — click any column header to sort (click again to reverse, once more to reset).",
            "<b>Filtering</b> — use the filters at the top and click Filter to narrow what you see.",
            "<b>Pagination</b> — choose how many rows per page (10 / 25 / 50 / 100) and move between pages.",
            "<b>Horizontal scrolling</b> — use the scrollbar at the top or bottom of the table to see more columns.",
            "<b>Sticky headers & fixed columns</b> — the header row and the Name/Date columns stay visible while you scroll.",
        ]),
        ("p", "To find a specific person quickly, type their name in the Employee Name filter and click Filter."),
    ]),
    ("10. Editing Existing Data", [
        ("p", "You can update or remove any observation you created."),
        ("h", "Editable vs locked"),
        ("bullets", [
            "<b>Editable</b>: all vigilance-entered fields (System Login/Logout, research/break hours, break blocks).",
            "<b>Locked</b>: Name, Email-id, Team, Date, Punch-In, Punch-Out, Total Hours.",
        ]),
        ("h", "Update an entry"),
        ("steps", [
            "Find the row in the table.",
            "Click the <b>pencil</b> icon to open the entry.",
            "Change the values you need.",
            "Click <b>Save Entry</b>.",
        ]),
        ("h", "Delete an entry"),
        ("steps", [
            "Find the row that has your entry.",
            "Click the <b>red trash</b> icon (it only appears once an entry exists).",
            "Confirm the deletion. Only that observation is removed — attendance and employee data are untouched.",
        ]),
    ]),
    ("11. Downloading / Exporting Data", [
        ("steps", [
            "Set any filters you want (dates, team, employee) and click Filter.",
            "Click <b>Export</b>.",
            "An Excel file downloads with exactly what is currently shown.",
        ]),
        ("note", "The export always respects your current filters — so filter first, then export to get precisely the records you need."),
    ]),
    ("12. Troubleshooting (FAQ)", [
        ("table", ["Question", "Solution"], FAQ),
    ]),
    ("13. Best Practices", [
        ("p", "Follow this simple, repeatable workflow every time for clean, error-free data:"),
        ("steps", [
            "<b>Apply filters</b> — set From/To dates and a Team/Department.",
            "<b>Download template</b> — get the pre-filled sheet.",
            "<b>Fill data</b> — enter your observations in the editable columns only.",
            "<b>Validate time format</b> — 24-hour clock times, HH:MM or HH:MM:SS durations.",
            "<b>Upload</b> — submit the completed sheet and read the result message.",
            "<b>Verify data</b> — filter to your dates and confirm the rows look correct.",
        ]),
        ("note", "Tip: keep date ranges short (a few days or one team at a time). Smaller batches are easier to fill correctly and quicker to fix if something needs a change."),
    ]),
]


# ===========================================================================
# PDF
# ===========================================================================
class GuideDoc(BaseDocTemplate):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._toc_entries = []

    def afterFlowable(self, flowable):
        if hasattr(flowable, "_toc_level"):
            self.notify("TOCEntry", (flowable._toc_level, flowable.getPlainText(), self.page, flowable._toc_key))
            self.canv.bookmarkPage(flowable._toc_key)
            self.canv.addOutlineEntry(flowable.getPlainText(), flowable._toc_key, level=flowable._toc_level)


def _styles():
    ss = getSampleStyleSheet()
    base = ss["BodyText"]
    body = ParagraphStyle("Body", parent=base, fontName="Helvetica", fontSize=10.5,
                          leading=15.5, alignment=TA_JUSTIFY, spaceAfter=6, textColor=colors.HexColor("#1F2937"))
    h1 = ParagraphStyle("H1", parent=base, fontName="Helvetica-Bold", fontSize=16,
                        textColor=NAVY, spaceBefore=6, spaceAfter=10, leading=20)
    h2 = ParagraphStyle("H2", parent=base, fontName="Helvetica-Bold", fontSize=12,
                        textColor=ACCENT, spaceBefore=8, spaceAfter=4, leading=16)
    note = ParagraphStyle("Note", parent=body, fontName="Helvetica-Oblique", fontSize=10,
                          leading=14, textColor=NAVY, alignment=TA_LEFT)
    li = ParagraphStyle("LI", parent=body, alignment=TA_LEFT, spaceAfter=3)
    cell = ParagraphStyle("Cell", parent=body, fontSize=9, leading=12.5, alignment=TA_LEFT, spaceAfter=0)
    cellb = ParagraphStyle("CellB", parent=cell, fontName="Helvetica-Bold", textColor=colors.white)
    toc_title = ParagraphStyle("TOCTitle", parent=h1)
    return dict(body=body, h1=h1, h2=h2, note=note, li=li, cell=cell, cellb=cellb, toc_title=toc_title)


def _heading(text, key, S):
    p = Paragraph(text, S["h1"])
    p._toc_level = 0
    p._toc_key = key
    return p


def _table(headers, rows, S, col_widths=None):
    data = [[Paragraph(str(h), S["cellb"]) for h in headers]]
    for r in rows:
        data.append([Paragraph(str(c), S["cell"]) for c in r])
    t = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D5DEE8")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            style.append(("BACKGROUND", (0, i), (-1, i), ROW_ALT))
    t.setStyle(TableStyle(style))
    return t


def _bullets(items, S):
    lis = [ListItem(Paragraph(x, S["li"]), leftIndent=6) for x in items]
    return ListFlowable(lis, bulletType="bullet", bulletColor=ACCENT, start="•", leftIndent=14)


def _steps(items, S):
    lis = [ListItem(Paragraph(x, S["li"]), leftIndent=6) for x in items]
    return ListFlowable(lis, bulletType="1", bulletFontName="Helvetica-Bold",
                        bulletColor=NAVY, leftIndent=16)


def _note(text, S):
    p = Paragraph(text, S["note"])
    box = Table([[p]], colWidths=[160 * mm])
    box.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LINEBEFORE", (0, 0), (0, -1), 3, EMERALD),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
    ]))
    return box


def _header_footer(canvas, doc):
    canvas.saveState()
    w, h = A4
    # Header band
    canvas.setFillColor(NAVY)
    canvas.rect(0, h - 16 * mm, w, 16 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(18 * mm, h - 10.5 * mm, f"{BRAND}")
    canvas.setFont("Helvetica", 8.5)
    canvas.drawRightString(w - 18 * mm, h - 10.5 * mm, "Vigilance Module — User Guide")
    # Footer
    canvas.setStrokeColor(colors.HexColor("#D5DEE8"))
    canvas.setLineWidth(0.5)
    canvas.line(18 * mm, 14 * mm, w - 18 * mm, 14 * mm)
    canvas.setFillColor(GREY)
    canvas.setFont("Helvetica", 8)
    canvas.drawString(18 * mm, 9 * mm, "Confidential — for Vigilance Team use only")
    canvas.drawRightString(w - 18 * mm, 9 * mm, f"Page {doc.page}")
    canvas.restoreState()


def _cover(canvas, doc):
    canvas.saveState()
    w, h = A4
    canvas.setFillColor(NAVY)
    canvas.rect(0, 0, w, h, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor("#13294F"))
    canvas.rect(0, h - 90 * mm, w, 90 * mm, fill=1, stroke=0)
    canvas.setFillColor(EMERALD)
    canvas.rect(0, h - 92 * mm, w, 2 * mm, fill=1, stroke=0)

    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 15)
    canvas.drawCentredString(w / 2, h - 40 * mm, BRAND)
    canvas.setFont("Helvetica", 10)
    canvas.setFillColor(colors.HexColor("#9FB3CC"))
    canvas.drawCentredString(w / 2, h - 47 * mm, "HRMS PLATFORM")

    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 34)
    canvas.drawCentredString(w / 2, h - 130 * mm, TITLE)
    canvas.setFont("Helvetica", 16)
    canvas.setFillColor(colors.HexColor("#C7D6EA"))
    canvas.drawCentredString(w / 2, h - 142 * mm, SUBTITLE)

    canvas.setFont("Helvetica-Oblique", 11)
    canvas.setFillColor(colors.HexColor("#9FB3CC"))
    canvas.drawCentredString(w / 2, 40 * mm, VERSION)
    canvas.drawCentredString(w / 2, 32 * mm, "Confidential — for Vigilance Team use only")
    canvas.restoreState()


def build_pdf():
    path = os.path.join(OUT_DIR, "Vigilance_Module_User_Guide.pdf")
    S = _styles()
    doc = GuideDoc(path, pagesize=A4,
                   leftMargin=18 * mm, rightMargin=18 * mm,
                   topMargin=22 * mm, bottomMargin=20 * mm,
                   title="Vigilance Module User Guide", author=BRAND)

    frame = Frame(doc.leftMargin, doc.bottomMargin,
                  doc.width, doc.height, id="main")
    cover_frame = Frame(0, 0, A4[0], A4[1], id="cover")
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=_cover),
        PageTemplate(id="Body", frames=[frame], onPage=_header_footer),
    ])

    story = [NextPageTemplate("Body"), PageBreak()]

    # Table of contents
    story.append(Paragraph("Table of Contents", S["h1"]))
    story.append(Spacer(1, 4))
    toc = TableOfContents()
    toc.levelStyles = [ParagraphStyle("TOC0", fontName="Helvetica", fontSize=11,
                                      leading=18, leftIndent=4, textColor=NAVY)]
    story.append(toc)
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    for idx, (title, blocks) in enumerate(SECTIONS):
        story.append(_heading(title, f"sec{idx}", S))
        for kind, *payload in blocks:
            if kind == "p":
                story.append(Paragraph(payload[0], S["body"]))
            elif kind == "h":
                story.append(Paragraph(payload[0], S["h2"]))
            elif kind == "steps":
                story.append(_steps(payload[0], S))
                story.append(Spacer(1, 4))
            elif kind == "bullets":
                story.append(_bullets(payload[0], S))
                story.append(Spacer(1, 4))
            elif kind == "note":
                story.append(Spacer(1, 2))
                story.append(_note(payload[0], S))
                story.append(Spacer(1, 6))
            elif kind == "table":
                story.append(Spacer(1, 2))
                story.append(_table(payload[0], payload[1], S))
                story.append(Spacer(1, 6))
        story.append(Spacer(1, 8))

    doc.multiBuild(story)
    return path


# ===========================================================================
# DOCX (fallback)
# ===========================================================================
def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _docx_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, hh in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(str(hh))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], NAVY_HEX)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val).replace("<b>", "").replace("</b>", ""))
            r.font.size = Pt(9)
            if ri % 2 == 1:
                _shade(cells[i], LIGHT_HEX)


def _clean(t):
    return t.replace("<b>", "").replace("</b>", "")


def build_docx():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(BRAND); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(TITLE + "\n" + SUBTITLE); r2.bold = True; r2.font.size = Pt(26); r2.font.color.rgb = NAVY_RGB
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(VERSION).italic = True
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Confidential — for Vigilance Team use only").font.size = Pt(9)
    doc.add_page_break()

    # TOC (static list)
    th = doc.add_paragraph(); thr = th.add_run("Table of Contents")
    thr.bold = True; thr.font.size = Pt(16); thr.font.color.rgb = NAVY_RGB
    for title, _ in SECTIONS:
        doc.add_paragraph(title, style='List Bullet')
    doc.add_page_break()

    for title, blocks in SECTIONS:
        hp = doc.add_paragraph(); hr = hp.add_run(title)
        hr.bold = True; hr.font.size = Pt(15); hr.font.color.rgb = NAVY_RGB
        for kind, *payload in blocks:
            if kind == "p":
                doc.add_paragraph(_clean(payload[0]))
            elif kind == "h":
                sp = doc.add_paragraph(); sr = sp.add_run(_clean(payload[0]))
                sr.bold = True; sr.font.size = Pt(12); sr.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            elif kind == "steps":
                for x in payload[0]:
                    doc.add_paragraph(_clean(x), style='List Number')
            elif kind == "bullets":
                for x in payload[0]:
                    doc.add_paragraph(_clean(x), style='List Bullet')
            elif kind == "note":
                np = doc.add_paragraph(); nr = np.add_run("Note: " + _clean(payload[0]))
                nr.italic = True; nr.font.color.rgb = NAVY_RGB
            elif kind == "table":
                _docx_table(doc, payload[0], payload[1])
                doc.add_paragraph()

    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run(f"{BRAND} — Vigilance Module User Guide • Confidential").font.size = Pt(8)

    path = os.path.join(OUT_DIR, "Vigilance_Module_User_Guide.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    pdf = build_pdf()
    docx = build_docx()
    print("PDF :", pdf, os.path.getsize(pdf), "bytes")
    print("DOCX:", docx, os.path.getsize(docx), "bytes")
